# -*- coding: utf-8 -*-
"""
基金产品方案：网页填表 -> 自动生成 Word(docx)

你反馈的问题：
- 生成的 Word 中，部分占位符替换后“未保留原格式”（例如：段落/编号/缩进/字体混排被破坏）。

本次更新的核心修复：
- 文本占位符替换改为 **XML 级别** 修改 w:t（支持占位符跨 run 拆分），并用 <w:br/> 表示换行，
  从而最大化保留模板原有的段落样式、编号列表、缩进、行距、混排格式等。
- 仅在需要插入图片时使用 python-docx 的 add_picture（图片字段不再把“图片路径”当作普通文本写入文档）。
- 修复原文件末尾多余字符导致的运行错误（文件末尾孤立的 `F`）。

目录约定：
- app.py
- templates/                # 存放 Word 模板（docx），按“产品类型”分子文件夹
    - 行业主题ETF/
        - xxx.docx
    - 宽基ETF/
        - yyy.docx
- static/index.html         # 前端页面
- output/                   # 生成文件输出目录（自动创建）

依赖：flask、python-docx（外加 Python 标准库）
"""

from __future__ import annotations

import datetime
import io
import random
import os
import re
import zipfile
from typing import Dict, Iterable, List, Optional, Tuple

from flask import Flask, abort, jsonify, request, send_from_directory, send_file
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm


# ============ 基础配置 ============
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
DOCX_TEMPLATE_ROOT = os.path.join(BASE_DIR, "templates")   # docx 模板根目录
PUBLIC_DIR = os.path.join(BASE_DIR, "public")              # 静态资源目录（Vercel 推荐）
PUBLIC_STATIC_DIR = os.path.join(PUBLIC_DIR, "static")    # public/static（浏览器访问 /static/...）

# Flask 不使用 Jinja 的 HTML templates（避免与 docx 模板目录 templates/ 冲突）
# 我们只从 static/index.html 提供页面
app = Flask(__name__, static_folder=PUBLIC_STATIC_DIR, static_url_path="/static", template_folder="web_templates")

# 匹配 {{占位符}} 的正则（占位符名可含中文）
PH_RE = re.compile(r"{{(.*?)}}")

# ============ 可选：下拉框配置（按需自行改） ============
# 规则：如果占位符 key 在 SELECT_OPTIONS 里，前端会渲染为下拉选择
SELECT_OPTIONS = {
    "基金类型": [
        "股票型基金",
        "混合型基金",
        "债券型基金",
        "货币市场基金",
        "基金中基金（FOF）",
        "QDII",
        "商品型基金",
        "REITs",
        "其他",
    ],
    "上市交易所": [
        "上海证券交易所",
        "深圳证券交易所",
        "北京证券交易所",
        "香港交易所",
        "其他",
    ],
    "ETF类型": [
        "跨市场ETF",
        "单市场ETF",
        "跨境ETF",
        "债券ETF",
        "商品ETF",
        "其他",
    ],
    "指数公司名称": [
        "中证指数有限公司",
        "深圳证券信息有限公司",
        "恒生指数有限公司",
        "其他",
    ],
    "指数公司简称": [
        "中证指数",
        "深证信息",
        "恒生指数",
        "其他",
    ],
}

# 如果占位符名包含这些关键词，默认用 textarea（长文本）
LONG_TEXT_KEYWORDS = ["风险", "揭示", "策略", "分析", "简介", "介绍", "说明", "情况", "内容"]


# ============ 通用工具函数 ============
def _safe_join_under_root(root: str, user_path: str) -> str:
    """防止路径穿越：只允许 root 目录下的子路径"""
    user_path = (user_path or "").strip()
    joined = os.path.abspath(os.path.join(root, user_path))
    root_abs = os.path.abspath(root)
    if not (joined == root_abs or joined.startswith(root_abs + os.sep)):
        raise ValueError("非法路径")
    return joined


def _sanitize_folder_name(name: str, max_len: int = 80) -> str:
    """
    清理 Windows / macOS 常见非法字符，防止路径穿越。
    """
    name = (name or "").strip()
    # 禁止路径分隔符与非法字符：\ / : * ? " < > |
    name = re.sub(r'[\\/:*?"<>|]+', "_", name)
    # 去掉首尾点号/空格（Windows 容易出问题）
    name = name.strip(" .")
    if not name:
        name = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    if len(name) > max_len:
        name = name[:max_len].rstrip(" .")
    return name


def _ensure_unique_subdir(parent_dir: str, base_name: str) -> Tuple[str, str]:
    """
    返回 (run_dir, final_name)，若同名存在则自动追加 _01 _02 ...
    """
    base_name = _sanitize_folder_name(base_name)
    candidate = base_name
    i = 1
    while os.path.exists(os.path.join(parent_dir, candidate)):
        candidate = f"{base_name}_{i:02d}"
        i += 1
    run_dir = os.path.join(parent_dir, candidate)
    os.makedirs(run_dir, exist_ok=True)
    return run_dir, candidate


def safe_filename(name: str) -> str:
    """Windows/macOS 通用文件名清洗"""
    name = (name or "").replace("\\", "_").replace("/", "_")
    name = re.sub(r'[<>:"|?*]', "_", name)
    name = re.sub(r"\s+", " ", name).strip()
    if not name:
        name = "output"
    return name[:180]


def list_docx_files(folder: str) -> List[str]:
    """列出 folder 下的 docx 模板（过滤临时文件 ~$.docx）"""
    if not os.path.isdir(folder):
        return []
    return sorted(
        [
            f
            for f in os.listdir(folder)
            if f.lower().endswith(".docx") and not f.startswith("~$")
        ]
    )


def get_product_folder(product_type: str) -> str:
    """
    product_type 对应 templates 下的子文件夹。
    特殊值 __root__ 表示 templates 根目录（兼容“模板直接放 templates/”的情况）
    """
    if not product_type or product_type == "__root__":
        return DOCX_TEMPLATE_ROOT
    return _safe_join_under_root(DOCX_TEMPLATE_ROOT, product_type)


def list_product_types() -> List[Dict[str, str]]:
    """
    自动扫描 templates/：
    - 子文件夹 = 产品类型
    - 如果 templates/ 根目录本身有 docx，也作为一个“默认(__root__)”产品类型
    """
    types: List[Dict[str, str]] = []

    if os.path.isdir(DOCX_TEMPLATE_ROOT):
        for name in sorted(os.listdir(DOCX_TEMPLATE_ROOT)):
            full = os.path.join(DOCX_TEMPLATE_ROOT, name)
            if os.path.isdir(full) and not name.startswith("."):
                if list_docx_files(full):
                    types.append({"id": name, "name": name})

    if list_docx_files(DOCX_TEMPLATE_ROOT):
        types.insert(0, {"id": "__root__", "name": "默认（templates根目录）"})

    return types


# ============ Word：扫描占位符（用于前端 schema） ============
def scan_placeholders_and_longflags(doc: Document) -> Tuple[set, set]:
    """
    扫描一个 doc，返回：
    - keys: 占位符集合
    - long_candidate: 哪些占位符“可能是长文本”（在某个段落中独占一行）
    """
    keys = set()
    long_candidate = set()

    def _scan_element(el):
        nonlocal keys, long_candidate
        for p in el.xpath(".//w:p"):
            text = "".join([t.text or "" for t in p.xpath(".//w:t")])
            text_stripped = text.strip()

            found = [m.group(1).strip() for m in PH_RE.finditer(text)]
            for k in found:
                keys.add(k)

            # “独占一行”的占位符 -> textarea 更合适（典型：风险揭示、投资策略大段）
            if len(found) == 1 and text_stripped == "{{%s}}" % found[0]:
                long_candidate.add(found[0])

    _scan_element(doc._part._element)

    for section in doc.sections:
        parts = [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ]
        for part in parts:
            if part is None:
                continue
            _scan_element(part._element)

    return keys, long_candidate


def build_schema(product_type: str) -> Dict:
    """
    根据产品类型，读取对应文件夹下所有 docx，自动扫描：
    - 模板文件列表
    - 占位符字段列表（并给出 text/textarea/select 的建议）
    """
    folder = get_product_folder(product_type)
    tpl_files = list_docx_files(folder)
    if not tpl_files:
        return {"product_type": product_type, "templates": [], "fields": []}

    all_keys = set()
    long_flags = set()

    for f in tpl_files:
        p = os.path.join(folder, f)
        doc = Document(p)
        keys, longs = scan_placeholders_and_longflags(doc)
        all_keys |= keys
        long_flags |= longs

    fields = []
    for key in sorted(all_keys):
        ftype = "text"
        if key in SELECT_OPTIONS:
            ftype = "select"
        elif key in long_flags or any(kw in key for kw in LONG_TEXT_KEYWORDS):
            ftype = "textarea"

        item = {"key": key, "label": key, "type": ftype}
        if ftype == "select":
            item["options"] = SELECT_OPTIONS.get(key, [])
        fields.append(item)

    return {
        "product_type": product_type,
        "templates": tpl_files,
        "fields": fields,
    }


# ============ ✅ 文本占位符替换：XML 级别（保留格式） ============
def _set_text_preserve_space(t_elem, text: str) -> None:
    """
    写入 w:t 文本；
    必要时加 xml:space='preserve' 防止首尾空格或连续空格被 Word 吞掉
    """
    t_elem.text = text
    if text and (text[0].isspace() or text[-1].isspace() or "  " in text):
        t_elem.set(qn("xml:space"), "preserve")
    else:
        attr = qn("xml:space")
        if attr in t_elem.attrib:
            del t_elem.attrib[attr]


def _insert_after(parent, ref_child, new_child):
    parent.insert(parent.index(ref_child) + 1, new_child)


def _append_text_with_breaks(first_t, text_to_insert: str) -> None:
    """
    在 first_t 所在 run 内追加文本：
    - 用户输入的 \\n 转成 <w:br/>（行内换行）
    - 不创建新段落 => 段落的行距/缩进/编号/样式最大程度保持模板原样
    """
    run = first_t.getparent()  # w:r
    segments = (text_to_insert or "").split("\n")

    existing = first_t.text or ""
    _set_text_preserve_space(first_t, existing + segments[0])

    last_elem = first_t
    for seg in segments[1:]:
        br = OxmlElement("w:br")
        _insert_after(run, last_elem, br)
        last_elem = br
        if seg != "":
            tnew = OxmlElement("w:t")
            _set_text_preserve_space(tnew, seg)
            _insert_after(run, last_elem, tnew)
            last_elem = tnew


def replace_placeholders_in_element(element, mapping: Dict[str, Optional[str]]) -> int:
    """
    在一个 XML element（文档主体/页眉页脚等）内替换所有 {{占位符}}
    - mapping[key] == None 表示“跳过替换”（保留原 {{key}}）
    返回：替换次数（用于调试）
    """
    replaced = 0

    for p in element.xpath(".//w:p"):
        guard = 0
        while True:
            guard += 1
            if guard > 500:
                # 防止异常模板导致死循环
                break

            t_elems = p.xpath(".//w:t")
            if not t_elems:
                break

            full_text = "".join([t.text or "" for t in t_elems])
            matches = list(PH_RE.finditer(full_text))
            if not matches:
                break

            # 找到第一个“mapping 中存在且 value != None”的占位符
            target = None
            for m in matches:
                key = (m.group(1) or "").strip()
                if key in mapping and mapping.get(key) is not None:
                    target = m
                    break

            if not target:
                break

            key = (target.group(1) or "").strip()
            replacement = str(mapping.get(key) or "")

            start, end = target.span()  # [start, end)

            # 建立 char index -> t_elem 的映射
            cum = []
            pos = 0
            for t in t_elems:
                l = len(t.text or "")
                cum.append((pos, pos + l))
                pos += l

            def locate(posi: int):
                for i, (s, e) in enumerate(cum):
                    if s <= posi < e:
                        return i, posi - s
                return None, None

            first_i, start_off = locate(start)
            last_i, _ = locate(end - 1)
            if first_i is None or last_i is None:
                break

            end_off = (end - 1 - cum[last_i][0]) + 1

            first_t = t_elems[first_i]
            last_t = t_elems[last_i]
            first_text = first_t.text or ""
            last_text = last_t.text or ""

            prefix = first_text[:start_off]
            suffix = last_text[end_off:]

            if first_i == last_i:
                # 占位符在同一个 w:t 内：prefix + replacement + suffix
                _set_text_preserve_space(first_t, prefix)
                _append_text_with_breaks(first_t, replacement + suffix)
            else:
                # 跨多个 w:t：在第一个 w:t 插入 replacement；最后一个 w:t 保留 suffix；中间清空
                _set_text_preserve_space(first_t, prefix)
                _append_text_with_breaks(first_t, replacement)

                for j in range(first_i + 1, last_i):
                    _set_text_preserve_space(t_elems[j], "")

                _set_text_preserve_space(last_t, suffix)

            replaced += 1

    return replaced


def replace_placeholders_in_doc(doc: Document, mapping: Dict[str, Optional[str]]) -> None:
    """替换正文 + 页眉页脚（含首页/奇偶页）"""
    replace_placeholders_in_element(doc._part._element, mapping)

    for section in doc.sections:
        parts = [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ]
        for part in parts:
            if part is None:
                continue
            replace_placeholders_in_element(part._element, mapping)


# ============ ✅ 图片占位符（可选） ============
def _normalize_newlines(s: str) -> str:
    return (s or "").replace("\r\n", "\n").replace("\r", "\n")


def _is_blank_only(s: str) -> bool:
    # strip() 会处理中文全角空格（\u3000）
    return (s or "").strip() == ""


def _iter_paragraphs_in_parent(parent) -> Iterable:
    """
    递归遍历：parent.paragraphs + parent.tables 中的 cell.paragraphs
    """
    for p in getattr(parent, "paragraphs", []):
        yield p
    for tbl in getattr(parent, "tables", []):
        for row in tbl.rows:
            for cell in row.cells:
                yield from _iter_paragraphs_in_parent(cell)


def _iter_all_paragraphs(doc: Document) -> Iterable:
    """
    遍历：
    - 正文 + 表格
    - 页眉页脚（含首页/奇偶页）
    """
    yield from _iter_paragraphs_in_parent(doc)
    for section in doc.sections:
        yield from _iter_paragraphs_in_parent(section.header)
        yield from _iter_paragraphs_in_parent(section.footer)
        try:
            yield from _iter_paragraphs_in_parent(section.first_page_header)
            yield from _iter_paragraphs_in_parent(section.first_page_footer)
            yield from _iter_paragraphs_in_parent(section.even_page_header)
            yield from _iter_paragraphs_in_parent(section.even_page_footer)
        except Exception:
            pass


def _resolve_local_path(base_dir: str, maybe_path: str) -> Optional[str]:
    if not maybe_path:
        return None
    p = str(maybe_path).strip()
    if not p:
        return None
    if os.path.isabs(p):
        return p
    p = p.replace("\\", "/").lstrip("/")
    # 兼容：presets/...（放在 static 下）
    if p.startswith("presets/"):
        return os.path.join(base_dir, "static", p)
    # 兼容：static/...
    if p.startswith("static/"):
        return os.path.join(base_dir, p)
    return os.path.join(base_dir, p)


def _clear_paragraph(paragraph) -> None:
    for r in paragraph.runs[::-1]:
        r._element.getparent().remove(r._element)


def _replace_image_in_paragraph(paragraph, placeholder: str, image_path: str, width_cm: Optional[float]) -> bool:
    """
    图片替换策略：
    - 若该段落除了占位符之外基本为空白：用图片“替换整段”（不保留原文字）
    - 否则：尽量在占位符位置插入图片
    """
    if placeholder not in paragraph.text:
        return False

    full = "".join(r.text for r in paragraph.runs) if paragraph.runs else paragraph.text
    idx = full.find(placeholder)
    prefix = full[:idx] if idx >= 0 else ""
    suffix = full[idx + len(placeholder) :] if idx >= 0 else ""

    if _is_blank_only(prefix + suffix):
        _clear_paragraph(paragraph)
        run = paragraph.add_run()
        if width_cm:
            run.add_picture(image_path, width=Cm(width_cm))
        else:
            run.add_picture(image_path)
        return True

    # 尝试在某个 run 内替换（最不破坏格式）
    for run in paragraph.runs:
        if placeholder in run.text:
            before, after = run.text.split(placeholder, 1)
            run.text = before

            pic_run = paragraph.add_run()
            if width_cm:
                pic_run.add_picture(image_path, width=Cm(width_cm))
            else:
                pic_run.add_picture(image_path)

            if after:
                paragraph.add_run(after)
            return True

    # 兜底：重建段落（可能损失混排，但保证能插入）
    before, _, after = full.partition(placeholder)
    _clear_paragraph(paragraph)
    paragraph.add_run(before)

    pic_run = paragraph.add_run()
    if width_cm:
        pic_run.add_picture(image_path, width=Cm(width_cm))
    else:
        pic_run.add_picture(image_path)

    paragraph.add_run(after)
    return True


def replace_images_in_doc(
    doc: Document,
    values_raw: Dict[str, str],
    base_dir: str,
    image_field_config: Optional[Dict[str, Dict]] = None,
) -> None:
    """
    image_field_config 形如：
    {
        "某图片字段": {"width_cm": 8}
    }
    values_raw["某图片字段"] = "static/xxx.png" 或 "presets/xxx.png" 或 绝对路径
    """
    image_field_config = image_field_config or {}
    if not image_field_config:
        return

    for key, cfg in image_field_config.items():
        placeholder = "{{" + key + "}}"
        img_path = _resolve_local_path(base_dir, values_raw.get(key))
        if not img_path or not os.path.exists(img_path):
            # 图片不存在：这里不做文本兜底，交给后续“文本替换”按 blank_unfilled 决策（空 or 保留占位符）
            continue

        for p in _iter_all_paragraphs(doc):
            if placeholder in p.text:
                _replace_image_in_paragraph(p, placeholder, img_path, cfg.get("width_cm"))


# ============ 生成辅助 ============
def build_docx_mapping(
    all_keys: List[str],
    values_raw: Dict[str, str],
    blank_unfilled: bool,
    image_keys: Optional[Iterable[str]] = None,
) -> Dict[str, Optional[str]]:
    """
    构造 XML 替换用 mapping：
    - blank_unfilled=True  ：所有字段都替换（未填 -> ""），避免残留 {{占位符}}
    - blank_unfilled=False ：仅替换“有内容”的字段，未填字段保持原 {{占位符}}
    - 图片字段（image_keys）：不把“图片路径”写入 docx（成功插图后占位符已消失；失败则按 blank_unfilled 决策）
    """
    image_keys_set = set(image_keys or [])
    mapping: Dict[str, Optional[str]] = {}

    for k in all_keys:
        if k in image_keys_set:
            mapping[k] = "" if blank_unfilled else None
            continue

        v = values_raw.get(k)
        v_str = "" if v is None else str(v)
        if blank_unfilled:
            mapping[k] = _normalize_newlines(v_str)
        else:
            mapping[k] = _normalize_newlines(v_str) if v_str.strip() != "" else None

    return mapping


def apply_placeholders_to_filename(filename: str, mapping: Dict[str, Optional[str]], blank_unfilled: bool = True) -> str:
    """把模板文件名里的 {{占位符}} 也替换掉，得到输出文件名"""

    def _rep(m):
        k = (m.group(1) or "").strip()
        if k not in mapping:
            return "" if blank_unfilled else m.group(0)

        v = mapping.get(k)
        if v is None or str(v).strip() == "":
            return "" if blank_unfilled else m.group(0)

        return str(v)

    name = PH_RE.sub(_rep, filename)

    # blank_unfilled=True 时，清掉仍残留的 {{...}}
    if blank_unfilled:
        name = re.sub(r"{{.*?}}", "", name)

    return safe_filename(name)


def create_zip_from_folder(folder: str, zip_path: str, arc_root_name: str = "") -> None:
    """
    把 folder 整个打包为 zip
    - arc_root_name 不为空时：zip 内会多一层顶层目录（更整洁，解压不散落）
    """
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for root, _, files in os.walk(folder):
            for fn in files:
                if fn.startswith("~$"):
                    continue
                full = os.path.join(root, fn)
                rel = os.path.relpath(full, folder)
                arc = os.path.join(arc_root_name, rel) if arc_root_name else rel
                z.write(full, arcname=arc)


# ============ 路由 ============
@app.route("/")
def index():
    # Vercel 会优先从 public/ 提供静态文件；本路由主要用于本地开发
    return send_from_directory(PUBLIC_DIR, "index.html")


@app.route("/api/product_types")
def api_product_types():
    if not os.path.isdir(DOCX_TEMPLATE_ROOT):
        return jsonify({"ok": False, "error": "templates 目录不存在"}), 400
    return jsonify({"ok": True, "data": list_product_types()})


@app.route("/api/schema")
def api_schema():
    product_type = request.args.get("product_type", "__root__")
    try:
        folder = get_product_folder(product_type)
    except Exception:
        return jsonify({"ok": False, "error": "product_type 非法"}), 400

    if not os.path.isdir(folder):
        return jsonify({"ok": False, "error": f"未找到产品类型目录：{product_type}"}), 404

    schema = build_schema(product_type)
    return jsonify({"ok": True, "data": schema})


@app.route("/api/generate", methods=["POST"])
def api_generate():
    """
    生成 docx 并直接返回 zip（二进制流）。

    说明：
    - 在 Vercel 这类无状态函数环境中，把文件写到项目目录再用 /download 二次下载不可靠。
    - 因此这里改为：一次请求内完成生成，并直接把 zip 作为响应返回。
    """
    payload = request.get_json(force=True, silent=False) or {}

    product_type = payload.get("product_type", "__root__")
    values_raw = payload.get("values", {}) or {}
    mode = payload.get("mode", "all")  # all / selected
    selected = payload.get("selected_templates", []) or []
    blank_unfilled = bool(payload.get("blank_unfilled", True))

    try:
        folder = get_product_folder(product_type)
    except Exception:
        return jsonify({"ok": False, "error": "product_type 非法"}), 400

    tpl_files = list_docx_files(folder)
    if not tpl_files:
        return jsonify({"ok": False, "error": f"该产品类型下没有 .docx 模板：{product_type}"}), 400

    if mode == "selected":
        tpl_files = [f for f in tpl_files if f in selected]
        if not tpl_files:
            return jsonify({"ok": False, "error": "未选择任何模板"}), 400

    # 读取 schema，用于得到 all_keys（确保 blank_unfilled=True 时也能把所有占位符清空）
    schema = build_schema(product_type)
    all_keys = [f["key"] for f in schema.get("fields", [])]

    # ========= 可选：图片字段配置（有图片占位符时启用） =========
    image_field_config = {
        "指数公司介绍图片": {"width_cm": 9.5},
    }
    image_keys = list(image_field_config.keys())

    mapping = build_docx_mapping(
        all_keys=all_keys,
        values_raw=values_raw,
        blank_unfilled=blank_unfilled,
        image_keys=image_keys,
    )

    # ========= zip 命名 =========
    output_folder = (payload.get("output_folder") or "").strip()
    if not output_folder:
        idx = str(values_raw.get("指数名称") or "").strip()
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_folder = f"{product_type}_{idx or '未命名'}_{ts}"

    run_name = _sanitize_folder_name(output_folder)
    # 避免浏览器多次下载同名覆盖，追加一个短随机后缀
    rand = ''.join(random.choice('ABCDEFGHJKLMNPQRSTUVWXYZ23456789') for _ in range(4))
    run_name = f"{run_name}_{rand}"

    zip_name = f"{run_name}.zip"

    # ========= 生成并打包（内存里完成） =========
    zip_buf = io.BytesIO()
    generated_files = []
    used_paths = set()

    try:
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as z:
            for tpl in tpl_files:
                tpl_path = os.path.join(folder, tpl)
                doc = Document(tpl_path)

                # 1) 图片占位符（可选）
                replace_images_in_doc(
                    doc,
                    values_raw=values_raw,
                    base_dir=BASE_DIR,
                    image_field_config=image_field_config,
                )

                # 2) 文本占位符：XML 级别替换，尽量保留模板格式
                replace_placeholders_in_doc(doc, mapping)

                out_name = apply_placeholders_to_filename(tpl, mapping, blank_unfilled=blank_unfilled)
                if not out_name.lower().endswith(".docx"):
                    out_name += ".docx"

                # 防重名：同名则自动加 (2)(3)...
                base_name = out_name
                counter = 2
                arc_path = f"{run_name}/{out_name}"
                while arc_path in used_paths:
                    stem, ext = os.path.splitext(base_name)
                    out_name = f"{stem}({counter}){ext}"
                    arc_path = f"{run_name}/{out_name}"
                    counter += 1

                used_paths.add(arc_path)

                bio = io.BytesIO()
                doc.save(bio)
                z.writestr(arc_path, bio.getvalue())

                generated_files.append(out_name)

    except Exception as e:
        return jsonify({"ok": False, "error": f"生成失败：{repr(e)}"}), 500

    zip_buf.seek(0)

    resp = send_file(
        zip_buf,
        mimetype="application/zip",
        as_attachment=True,
        download_name=zip_name,
        max_age=0,
    )
    # 轻量元信息（可选）
    resp.headers["X-Job-Id"] = run_name
    resp.headers["X-Generated-Count"] = str(len(generated_files))
    return resp

# 注意：Vercel 的无状态运行环境不适合把生成文件落盘后再二次下载。
# 当前版本 /api/generate 会直接返回 zip 文件（前端用 blob 下载）。


if __name__ == "__main__":
    # 本地运行：python app.py
    app.run(host="0.0.0.0", port=5000, debug=True)